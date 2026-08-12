from __future__ import annotations

import io
import json
import logging
import os
import re
from typing import Any

import PyPDF2
from docx import Document
from groq import Groq
from pydantic import ValidationError

from schemas import (
    BloomsLevel,
    Difficulty,
    ExamGenerationResponse,
    GenerateExamRequest,
    GradeSubjectiveRequest,
    GradeSubjectiveResponse,
    QuestionSchema,
)

logger = logging.getLogger("examora.ai")

DEFAULT_MODEL = "llama3-70b-8192"
MAX_GENERATION_ATTEMPTS = 3
MAX_DOCUMENT_BYTES = 10 * 1024 * 1024
MAX_DOCUMENT_CHARS_FOR_LLM = 60_000
SUPPORTED_DOCUMENT_EXTENSIONS = frozenset({"pdf", "docx", "txt", "md"})

# Only the first N pages of a PDF are parsed/OCR'd. Rendering a whole
# hundreds-page PDF page-by-page (or at once) would exhaust container RAM.
MAX_DOCUMENT_PAGES = 60

# Below this many characters, PyPDF2 extraction is treated as failed
# (scanned/image-only PDF) and the OCR fallback kicks in.
MIN_PDF_TEXT_CHARS = 20
OCR_DPI = 200

MIXED_DISTRIBUTION = {"easy": 0.3, "medium": 0.5, "hard": 0.2}

GENERATION_SYSTEM_PROMPT = """You are the Examora AI Exam Generator, an expert educational content creator.
Generate a complete exam question bank as a SINGLE raw JSON object.

OUTPUT FORMAT (STRICT):
Respond with ONLY a valid raw JSON object. No markdown, no code fences, no commentary, no text before or after the JSON.
The JSON MUST exactly match this schema:

{
  "topic": "string",
  "subtopics": ["string"],
  "difficulty": "easy" | "medium" | "hard" | "mixed",
  "blooms_level": "remember" | "understand" | "apply" | "analyze" | "evaluate" | "create" | null,
  "questions": [
    {
      "id": "q1",
      "question_text": "string",
      "type": "mcq_single" | "mcq_multi" | "true_false" | "short_answer" | "long_answer" | "fill_blank" | "dropdown",
      "options": ["string", "string", "string", "string"],
      "correct_answer": "string",
      "explanation": "string",
      "marks": 2,
      "difficulty": "easy" | "medium" | "hard",
      "blooms_level": "remember" | "understand" | "apply" | "analyze" | "evaluate" | "create",
      "confidence": 0.95
    }
  ]
}

HARD RULES:
1. Generate EXACTLY the requested number of questions. Never fewer, never more.
2. Question type selection:
   - mcq_single: 4 distinct options, correct_answer is exactly one option.
   - mcq_multi: 4-6 distinct options, correct_answer lists every correct option joined by "; " (at least 2).
   - true_false: options are exactly ["True", "False"], correct_answer is "True" or "False".
   - short_answer / long_answer: options is an empty array, correct_answer is a sample model answer.
   - fill_blank: question_text contains a blank marked with "___", options is an empty array.
3. NEVER include meta options such as "All of the above", "None of the above", "All are correct", "None are correct", or any equivalent phrasing.
4. Every single question MUST include a non-empty, pedagogically useful "explanation" field.
5. Options within one question MUST be mutually exclusive and distinct.
6. For choice types, "correct_answer" MUST exactly match one of the "options" strings (or a ";"-joined subset for mcq_multi).
7. The exam must be factually accurate, curriculum-appropriate, and aligned to the given difficulty and Bloom's level.
8. "confidence" is a 0.0-1.0 estimate of how certain you are the question is factually correct and well-formed.
9. Question ids must be "q1"..."qN" in order. Prefer standard MCQ/dropdown questions for most of the exam; use open-ended types sparingly."""

PARSE_DOCUMENT_SYSTEM_PROMPT = """You are the Examora AI Document Parser, an expert at converting exam papers into structured question banks.
You receive the extracted text of a document (PDF, DOCX or plain text) containing exam questions.
Extract every question you can identify and return a SINGLE raw JSON object.

OUTPUT FORMAT (STRICT):
Respond with ONLY a valid raw JSON object. No markdown, no code fences, no commentary, no text before or after the JSON.
The JSON MUST exactly match this schema:

{
  "topic": "best guess of the exam subject/topic from the document",
  "subtopics": ["string"],
  "difficulty": "easy" | "medium" | "hard" | "mixed",
  "blooms_level": "remember" | "understand" | "apply" | "analyze" | "evaluate" | "create" | null,
  "questions": [
    {
      "id": "q1",
      "question_text": "string (cleaned, keep the original wording)",
      "type": "mcq_single" | "mcq_multi" | "true_false" | "short_answer" | "long_answer" | "fill_blank" | "dropdown" | "date" | "file_upload",
      "options": ["string"],
      "correct_answer": "string",
      "explanation": "string",
      "marks": 2,
      "difficulty": "easy" | "medium" | "hard",
      "blooms_level": "remember" | "understand" | "apply" | "analyze" | "evaluate" | "create",
      "confidence": 0.85
    }
  ]
}

HARD RULES:
1. Auto-detect each question's type from its original structure (multiple choice, true/false, fill in the blank, short answer, long answer, etc.).
2. If a choice-based question in the document has missing or partial options, GENERATE plausible distractors to complete it.
3. If a correct answer is present in the document, transcribe it; otherwise provide your best scholarly answer.
4. NEVER invent meta options such as "All of the above", "None of the above", "All are correct", "None are correct".
5. Every question MUST include a non-empty "explanation".
6. "confidence" (0.0-1.0) reflects how faithfully the question was extracted and how reliable the answer is.
7. Preserve the original numbering semantics but reassign ids as "q1"..."qN" in document order.
8. If the document contains no identifiable questions, return "questions": []. Do not fabricate content that is not in the document."""


def compute_target_counts(count: int) -> dict[str, int]:
    raw = {difficulty: count * share for difficulty, share in MIXED_DISTRIBUTION.items()}
    counts = {difficulty: int(value) for difficulty, value in raw.items()}
    remaining = count - sum(counts.values())
    for difficulty in sorted(raw, key=lambda d: raw[d] - counts[d], reverse=True):
        if remaining <= 0:
            break
        counts[difficulty] += 1
        remaining -= 1
    return counts


GRADE_SUBJECTIVE_SYSTEM_PROMPT = """You are the Examora AI Subjective Grader, an expert examiner who marks long/short-answer responses against a rubric.

OUTPUT FORMAT (STRICT):
Respond with ONLY a valid raw JSON object. No markdown, no code fences, no commentary.
The JSON MUST exactly match this schema:
{
  "marks_awarded": 3.5,
  "confidence": 0.85,
  "feedback": "Concise, constructive feedback for the student (1-3 sentences)."
}

GRADING RUBRIC (apply strictly):
1. Read the question, the student's answer, and the maximum marks.
2. Award marks proportionally for the correct ideas, definitions, steps and keywords the answer contains. Partial credit is encouraged for partially correct answers.
3. Deduct marks for unsupported claims, contradictions or off-topic content. NEVER award marks for an empty or blank answer.
4. marks_awarded MUST be a non-negative number between 0 and max_marks inclusive. Prefer halves (e.g. 0.5, 1.5) for partial credit.
5. "confidence" (0.0-1.0) reflects how certain you are that the marks are fair and the answer was unambiguous. Subjective, vague, or borderline answers must score BELOW 0.6.
6. "feedback" must be constructive, specific and in the same language as the answer.
7. If the answer is unrelated to the question, blank, or just repeated question text, award 0 marks and explain why."""


class AIServiceError(Exception):
    def __init__(self, message: str, code: str = "ai_error", status_code: int = 500):
        super().__init__(message)
        self.message = message
        self.code = code
        self.status_code = status_code


class AIQuestionService:
    def __init__(self, api_key: str | None = None, model: str | None = None):
        resolved_key = api_key or os.environ.get("GROQ_API_KEY")
        self.model = model or os.environ.get("GROQ_MODEL") or DEFAULT_MODEL
        self._client = Groq(api_key=resolved_key) if resolved_key else None

    @property
    def available(self) -> bool:
        return self._client is not None

    def generate_exam(self, request: GenerateExamRequest) -> ExamGenerationResponse:
        if not self._client:
            raise AIServiceError("GROQ_API_KEY is not configured", code="service_unavailable", status_code=503)

        user_prompt = self._build_generation_user_prompt(request)
        attempts = MAX_GENERATION_ATTEMPTS
        last_errors: list[str] = []

        while attempts > 0:
            corrective = f"\n\nCORRECTION (previous attempt rejected): {'. '.join(last_errors)}\nReturn the corrected JSON object only." if last_errors else ""
            payload = self._complete_json(user_prompt + corrective, GENERATION_SYSTEM_PROMPT, max_tokens=8192)
            try:
                response = ExamGenerationResponse.model_validate(payload)
                if len(response.questions) != request.questionCount:
                    raise ValidationError.from_exception_data(
                        "ExamGenerationResponse",
                        [{"type": "value_error", "loc": ("questions",), "msg": f"expected {request.questionCount} questions, got {len(response.questions)}", "input": payload}],
                    )
                response = self._enforce_mixed_distribution(response)
                response.generation_metadata.update(
                    {"model": self.model, "distribution_enforced": request.difficulty == Difficulty.MIXED, "attempts": MAX_GENERATION_ATTEMPTS - attempts + 1}
                )
                return response
            except ValidationError as exc:
                last_errors = [error["msg"] for error in exc.errors()][:5]
                logger.warning("Groq output failed validation (attempt %d): %s", MAX_GENERATION_ATTEMPTS - attempts + 1, last_errors)
                attempts -= 1

        raise AIServiceError("Groq repeatedly returned output that failed schema validation", code="invalid_response", status_code=502)

    def parse_document(self, filename: str, content: bytes, content_type: str | None = None) -> ExamGenerationResponse:
        if not self._client:
            raise AIServiceError("GROQ_API_KEY is not configured", code="service_unavailable", status_code=503)
        if not content:
            raise AIServiceError("uploaded file is empty", code="empty_document", status_code=422)
        if len(content) > MAX_DOCUMENT_BYTES:
            raise AIServiceError("file exceeds the 10 MB limit", code="file_too_large", status_code=413)

        text = self._extract_text(filename, content)
        if len(text.strip()) < 10:
            raise AIServiceError("no readable text could be extracted from the document", code="no_readable_text", status_code=422)

        user_prompt = (
            f"Extracted document text (exam paper, {len(text)} characters):\n\n"
            f"---DOCUMENT START---\n{text[:MAX_DOCUMENT_CHARS_FOR_LLM]}\n---DOCUMENT END---\n\n"
            "Extract the questions now. If none exist, return {\"questions\": []}."
        )

        attempts = MAX_GENERATION_ATTEMPTS
        last_errors: list[str] = []
        while attempts > 0:
            corrective = f"\n\nCORRECTION (previous attempt rejected): {'. '.join(last_errors)}\nReturn the corrected JSON object only." if last_errors else ""
            payload = self._complete_json(user_prompt + corrective, PARSE_DOCUMENT_SYSTEM_PROMPT, max_tokens=8192)
            try:
                response = ExamGenerationResponse.model_validate(payload)
            except ValidationError as exc:
                last_errors = [error["msg"] for error in exc.errors()][:5]
                logger.warning("Document parse failed validation (attempt %d): %s", MAX_GENERATION_ATTEMPTS - attempts + 1, last_errors)
                attempts -= 1
                continue
            if not response.questions:
                raise AIServiceError("no questions were identified in the document", code="no_questions_found", status_code=422)
            response.generation_metadata.update({"model": self.model, "source": "parse_document", "document": filename, "attempts": MAX_GENERATION_ATTEMPTS - attempts + 1})
            return response

        raise AIServiceError("Groq repeatedly returned output that failed schema validation", code="invalid_response", status_code=502)

    def grade_subjective(self, request: GradeSubjectiveRequest) -> GradeSubjectiveResponse:
        """Grades a short/long-answer response against the question with Groq.

        Returns marks (0..max_marks), a confidence score (0..1) and feedback.
        Low-confidence (< 0.6) results are flagged for manual review by the
        Express grading service.
        """
        if not self._client:
            raise AIServiceError("GROQ_API_KEY is not configured", code="service_unavailable", status_code=503)

        user_prompt = (
            f"Question:\n{request.question_text}\n\n"
            f"Maximum marks: {request.max_marks}\n\n"
            f"Student answer:\n{request.student_answer}\n\n"
            "Grade this answer now and return the JSON object only."
        )

        attempts = MAX_GENERATION_ATTEMPTS
        last_errors: list[str] = []
        while attempts > 0:
            corrective = f"\n\nCORRECTION (previous attempt rejected): {'. '.join(last_errors)}\nReturn the corrected JSON object only." if last_errors else ""
            payload = self._complete_json(
                user_prompt + corrective,
                GRADE_SUBJECTIVE_SYSTEM_PROMPT,
                max_tokens=1024,
            )
            try:
                response = GradeSubjectiveResponse.model_validate(payload)
            except ValidationError as exc:
                last_errors = [error["msg"] for error in exc.errors()][:5]
                logger.warning("Groq grading failed validation (attempt %d): %s", MAX_GENERATION_ATTEMPTS - attempts + 1, last_errors)
                attempts -= 1
                continue

            response.marks_awarded = min(float(response.marks_awarded), float(request.max_marks))
            response.confidence = max(0.0, min(1.0, float(response.confidence)))
            response.model = self.model
            return response

        raise AIServiceError("Groq repeatedly returned output that failed validation", code="invalid_response", status_code=502)

    def _build_generation_user_prompt(self, request: GenerateExamRequest) -> str:
        subtopics = ", ".join(request.subtopics) if request.subtopics else "any relevant subtopics"
        lines = [
            f"Topic: {request.topic}",
            f"Subtopics: {subtopics}",
            f"Question count: {request.questionCount}",
            f"Difficulty: {request.difficulty.value}",
        ]
        if request.difficulty == Difficulty.MIXED:
            counts = compute_target_counts(request.questionCount)
            lines.append(f"MIXED distribution MUST be exactly: {counts['easy']} easy, {counts['medium']} medium, {counts['hard']} hard.")
        lines.append(f"Bloom's taxonomy level: {request.bloomsLevel.value if request.bloomsLevel else 'any (vary levels across questions)'}")
        lines.append("Return the complete exam JSON object now.")
        return "\n".join(lines)

    def _complete_json(self, user_prompt: str, system_prompt: str, max_tokens: int) -> dict[str, Any]:
        try:
            completion = self._client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.4,
                max_tokens=max_tokens,
                response_format={"type": "json_object"},
            )
        except Exception as exc:
            status = getattr(exc, "status_code", None)
            if status == 429:
                raise AIServiceError("Groq rate limit exceeded, retry later", code="rate_limited", status_code=429) from exc
            raise AIServiceError(f"Groq API call failed: {exc}", code="groq_error", status_code=502) from exc

        raw = completion.choices[0].message.content or ""
        return self._extract_json(raw)

    @staticmethod
    def _extract_json(raw: str) -> dict[str, Any]:
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
            cleaned = re.sub(r"\s*```$", "", cleaned)
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise AIServiceError("Groq response contained no JSON object", code="invalid_response", status_code=502)
        try:
            return json.loads(cleaned[start : end + 1])
        except json.JSONDecodeError as exc:
            raise AIServiceError(f"Groq returned malformed JSON: {exc}", code="invalid_response", status_code=502) from exc

    @staticmethod
    def _extract_text(filename: str, content: bytes) -> str:
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
            raise AIServiceError(
                f"unsupported file type '{extension}', expected one of: {', '.join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))}",
                code="unsupported_format",
                status_code=415,
            )

        if extension == "pdf":
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages[:MAX_DOCUMENT_PAGES]]
            pages_skipped = max(0, len(reader.pages) - MAX_DOCUMENT_PAGES)
            if pages_skipped:
                logger.info("PDF has %d pages — only the first %d were parsed", len(reader.pages), MAX_DOCUMENT_PAGES)
            text = "\n\n".join(pages).strip()
            # Scanned PDFs yield empty (or near-empty) text from PyPDF2 —
            # fall back to OCR (pdf2image + pytesseract) in that case.
            if len(text) < MIN_PDF_TEXT_CHARS:
                ocr_text = AIQuestionService._ocr_pdf(content)
                if ocr_text.strip():
                    logger.info("Scanned PDF detected — OCR extracted %d characters", len(ocr_text))
                    return ocr_text
                logger.warning("OCR fallback produced no text for the scanned PDF")
            return text

        if extension == "docx":
            document = Document(io.BytesIO(content))
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(part for part in parts if part.strip())

        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise AIServiceError("could not decode the text file", code="no_readable_text", status_code=422)

    @staticmethod
    def _ocr_pdf(content: bytes, max_pages: int = MAX_DOCUMENT_PAGES) -> str:
        """OCR a scanned PDF one page at a time.

        Rendering only one page per convert_from_bytes call keeps peak memory
        bounded (a full-document render of a many-page PDF would spike RAM
        and OOM the container). Lazily imports pdf2image + pytesseract so the
        rest of the service keeps working when the system OCR dependencies
        (tesseract binary, poppler) are not installed.
        """
        try:
            from pdf2image import convert_from_bytes  # type: ignore[import-not-found]
            import pytesseract  # type: ignore[import-not-found]
        except ImportError as exc:
            logger.warning(
                "OCR fallback unavailable (%s) — install pytesseract + pdf2image and "
                "the tesseract/poppler system binaries to parse scanned PDFs",
                exc,
            )
            return ""

        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            total_pages = len(reader.pages)
        except Exception as exc:
            logger.warning("failed to count PDF pages for OCR: %s", exc)
            total_pages = max_pages

        pages_to_scan = min(total_pages, max_pages)
        if total_pages > max_pages:
            logger.info("OCR: scanning the first %d of %d pages", max_pages, total_pages)

        ocr_parts: list[str] = []
        for page_no in range(1, pages_to_scan + 1):
            try:
                images = convert_from_bytes(
                    content,
                    dpi=OCR_DPI,
                    first_page=page_no,
                    last_page=page_no,
                )
            except Exception as exc:
                logger.warning("pdf2image failed to render page %d: %s", page_no, exc)
                continue

            for image in images:
                try:
                    page_text = pytesseract.image_to_string(image).strip()
                except Exception as exc:
                    logger.warning("pytesseract failed on page %d: %s", page_no, exc)
                    continue
                if page_text:
                    ocr_parts.append(page_text)
        return "\n\n".join(ocr_parts)

    @staticmethod
    def _enforce_mixed_distribution(response: ExamGenerationResponse) -> ExamGenerationResponse:
        if response.difficulty != Difficulty.MIXED:
            return response

        targets = compute_target_counts(len(response.questions))
        current: dict[str, int] = {"easy": 0, "medium": 0, "hard": 0}
        for question in response.questions:
            current[question.difficulty.value] += 1

        questions = list(response.questions)
        questions.sort(key=lambda q: q.confidence)
        for question in questions:
            if all(current[d] == targets[d] for d in targets):
                break
            surplus = next((d for d in targets if current[d] > targets[d]), None)
            deficit = next((d for d in targets if current[d] < targets[d]), None)
            if surplus is None or deficit is None:
                break
            if question.difficulty.value == surplus:
                question.difficulty = Difficulty(deficit)
                current[surplus] -= 1
                current[deficit] += 1

        questions.sort(key=lambda q: int(re.sub(r"\D", "", q.id) or 0))
        response.questions = questions
        response.generation_metadata["difficulty_distribution"] = current
        return response
